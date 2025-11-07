from docx import Document

# Create a new document
document = Document() 
document.add_heading('My Document', 0)
document.add_paragraph('This is a paragraph created with python-docx.')

# Save the document
document.save('my_new_file.docx')
print(3)
