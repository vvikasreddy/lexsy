import os
import re
import tempfile
from dataclasses import dataclass
from typing import List, Optional, Tuple


DOCX_ENABLED = True
try:
    from docx import Document  # type: ignore
except Exception:
    DOCX_ENABLED = False


BLANK_PATTERNS = [
    re.compile(r"\$\[[^\]]*\]"),         # $[____] or $[placeholder]
    re.compile(r"(?<!\$)\[[^\]]+\]"),    # [something] that is NOT preceded by a dollar
    re.compile(r"_{2,}"),                   # __ or longer underscores (2+)
]


@dataclass
class BlankMatch:
    paragraph_index: int
    start: int
    end: int
    text: str
    label: str
    kind: str  # 'underscore' | 'bracket' | 'dollar_bracket'
    underscore_len: int
    before_words: List[str]
    after_words: List[str]

    @property
    def key(self) -> str:
        return f"p{self.paragraph_index}:{self.start}-{self.end}:{self.text}"


def _extract_words(s: str) -> List[str]:
    return re.findall(r"[A-Za-z0-9']+", s)


def _label_for_blank(s: str) -> str:
    # If [label] or $[label], return inner text; for underscores, generic label
    m = re.match(r"^\$?\[([^\]]*)\]$", s)
    if m:
        inner = m.group(1).strip() or "blank"
        return inner
    # Treat underscore-only blanks as generic label
    if re.fullmatch(r"_{2,}", s):
        return "blank"
    # If the text contains at least two underscore groups (e.g., "__/__/____"), treat as a generic blank
    if re.search(r"(?:(?<!_)_{2,}(?!_).*){2,}", s):
        return "blank"
    return s.strip("[]$") or "blank"


def _classify_blank(s: str) -> tuple[str, int]:
    if s.startswith("$[") and s.endswith("]"):
        return ("dollar_bracket", 0)
    if s.startswith("[") and s.endswith("]"):
        return ("bracket", 0)
    # Classify a plain underscore run (2 or more)
    m = re.fullmatch(r"(_{2,})", s)
    if m:
        return ("underscore", len(m.group(1)))
    return ("unknown", 0)


def find_blanks_in_docx(path: str, ctx_words: int = 15) -> List[BlankMatch]:
    if not (DOCX_ENABLED and os.path.exists(path)):
        return []
    doc = Document(path)
    matches: List[BlankMatch] = []
    for pi, p in enumerate(doc.paragraphs):
        text = p.text or ""
        if not text.strip():
            continue
        spans: List[Tuple[int, int, str]] = []
        for pat in BLANK_PATTERNS:
            for m in pat.finditer(text):
                spans.append((m.start(), m.end(), m.group(0)))
        # de-dup overlapping by start/end
        # De-dup exact and nested overlaps: keep longest / outermost span
        spans = sorted(set(spans))
        filtered: List[Tuple[int, int, str]] = []
        for s in spans:
            s_start, s_end, _ = s
            contained = False
            for t in spans:
                if t is s:
                    continue
                t_start, t_end, _ = t
                if t_start <= s_start and t_end >= s_end:
                    if (t_end - t_start) >= (s_end - s_start):
                        contained = True
                        break
            if not contained:
                filtered.append(s)
        spans = filtered
        for start, end, raw in spans:
            before = text[:start]
            after = text[end:]
            before_words = _extract_words(before)[-ctx_words:]
            after_words = _extract_words(after)[:ctx_words]
            kind, ulen = _classify_blank(raw)
            matches.append(
                BlankMatch(
                    paragraph_index=pi,
                    start=start,
                    end=end,
                    text=raw,
                    label=_label_for_blank(raw),
                    kind=kind,
                    underscore_len=ulen,
                    before_words=before_words,
                    after_words=after_words,
                )
            )
        # Also scan table cell paragraphs
    try:
        for ti, tbl in enumerate(doc.tables):
            for ri, row in enumerate(tbl.rows):
                for ci, cell in enumerate(row.cells):
                    for pi, p in enumerate(cell.paragraphs):
                        text = p.text or ""
                        if not text.strip():
                            continue
                        spans: List[Tuple[int, int, str]] = []
                        for pat in BLANK_PATTERNS:
                            for m in pat.finditer(text):
                                spans.append((m.start(), m.end(), m.group(0)))
                        spans = sorted(set(spans))
                        filtered: List[Tuple[int, int, str]] = []
                        for s in spans:
                            s_start, s_end, _ = s
                            contained = False
                            for t in spans:
                                if t is s:
                                    continue
                                t_start, t_end, _ = t
                                if t_start <= s_start and t_end >= s_end:
                                    if (t_end - t_start) >= (s_end - s_start):
                                        contained = True
                                        break
                            if not contained:
                                filtered.append(s)
                        spans = filtered
                        for start, end, raw in spans:
                            before = text[:start]
                            after = text[end:]
                            before_words = _extract_words(before)[-ctx_words:]
                            after_words = _extract_words(after)[:ctx_words]
                            kind, ulen = _classify_blank(raw)
                            matches.append(
                                BlankMatch(
                                    paragraph_index=-(len(matches) + 1),
                                    start=start,
                                    end=end,
                                    text=raw,
                                    label=_label_for_blank(raw),
                                    kind=kind,
                                    underscore_len=ulen,
                                    before_words=before_words,
                                    after_words=after_words,
                                    table_path=(ti, ri, ci, pi),
                                )
                            )
    except Exception:
        pass
    return matches


def replace_one_blank(
    path_in: str,
    paragraph_index: int,
    start: int,
    end: int,
    replacement: str,
    path_out: Optional[str] = None,
    table_path: Optional[Tuple[int, int, int, int]] = None,
) -> str:
    if not DOCX_ENABLED:
        raise RuntimeError("python-docx not installed")
    if not os.path.exists(path_in):
        raise FileNotFoundError(path_in)

    doc = Document(path_in)

    # Locate the paragraph either in a table cell or in the body
    if table_path is not None:
        ti, ri, ci, pi = table_path
        p = doc.tables[ti].rows[ri].cells[ci].paragraphs[pi]
    else:
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            raise IndexError("paragraph index out of range")
        p = doc.paragraphs[paragraph_index]

    text = p.text or ""
    new_text = text[:start] + replacement + text[end:]

    # Note: setting p.text resets run-level formatting for this paragraph
    p.text = new_text

    if not path_out:
        path_out = path_in
    doc.save(path_out)
    return path_out
